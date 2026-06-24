"""
Document Processing Utilities

This module provides utilities for extracting text from various document formats
including PDF, Word documents, and plain text files.

Enhanced with Docling support for better table/figure extraction.
"""

import io

try:
    from docling.document_converter import DocumentConverter
    DOCKING_AVAILABLE = True
except ImportError:
    DOCKING_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from ..config.logfire_config import get_logger, logfire

logger = get_logger(__name__)


def extract_text_from_document(file_content: bytes, filename: str, content_type: str) -> str:
    """Extract text from various document formats."""
    try:
        # PDF files
        if content_type == "application/pdf" or filename.lower().endswith(".pdf"):
            return extract_text_from_pdf(file_content)

        # Word documents
        elif content_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ] or filename.lower().endswith((".docx", ".doc")):
            return extract_text_from_docx(file_content)

        # Text files
        elif content_type.startswith("text/") or filename.lower().endswith((
            ".txt", ".md", ".markdown", ".rst",
        )):
            return file_content.decode("utf-8", errors="ignore")

        else:
            raise ValueError(f"Unsupported file format: {content_type} ({filename})")

    except Exception as e:
        logfire.error(
            "Document text extraction failed",
            filename=filename,
            content_type=content_type,
            error=str(e),
        )
        raise Exception(f"Failed to extract text from {filename}: {str(e)}") from e


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF using Docling (preferred) or fallback to PyPDF2/pdfplumber.

    Docling provides:
    - Better table extraction with structure preserved
    - Figure detection
    - Layout-aware parsing
    - 30x faster than OCR-based methods
    """
    # Try Docling first (best quality)
    if DOCKING_AVAILABLE:
        try:
            converter = DocumentConverter()
            result = converter.convert_stream(io.BytesIO(file_content))

            # Export to markdown (preserves structure better)
            markdown_content = result.document.export_to_markdown()

            if markdown_content and len(markdown_content.strip()) > 100:
                logger.debug("Docling extracted PDF content successfully")
                return markdown_content

            logger.warning("Docling extracted minimal content, falling back")
        except Exception as e:
            logger.warning(f"Docling extraction failed: {e}, falling back")

    # Fallback to pdfplumber (good for complex layouts)
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                text_parts = []
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
                return "\n\n".join(text_parts)
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}")

    # Final fallback to PyPDF2 (basic but reliable)
    if PYPDF2_AVAILABLE:
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
            return "\n\n".join(text_parts)
        except Exception as e:
            raise Exception(f"PyPDF2 failed to extract text: {str(e)}") from e

    raise Exception(
        "No PDF processing libraries available. Install: pip install docling pdfplumber PyPDF2"
    )


def extract_document_structure(file_content: bytes, filename: str) -> dict:
    """
    Extract structured information using Docling.

    Returns tables, figures, and structured text content.
    """
    if not DOCKING_AVAILABLE:
        return {"text": extract_text_from_pdf(file_content), "tables": [], "figures": []}

    try:
        converter = DocumentConverter()
        result = converter.convert_stream(io.BytesIO(file_content))

        return {
            "text": result.document.export_to_text(),
            "markdown": result.document.export_to_markdown(),
            "json": result.document.export_to_dict(),
            "tables": [tbl.export_to_dataframe().to_dict() for tbl in result.document.tables],
            "figures": [fig.image_base64 for fig in result.document.figures],
        }
    except Exception as e:
        logger.error(f"Structured extraction failed: {e}")
        return {
            "text": extract_text_from_pdf(file_content),
            "tables": [],
            "figures": [],
            "error": str(e)
        }


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from Word documents (.docx)."""
    if not DOCX_AVAILABLE:
        raise Exception("python-docx library not available. Please install python-docx.")

    try:
        doc = DocxDocument(io.BytesIO(file_content))
        text_content = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_content.append(" | ".join(row_text))

        if not text_content:
            raise Exception("No text content found in document")

        return "\n\n".join(text_content)

    except Exception as e:
        raise Exception(f"Failed to extract text from Word document: {str(e)}") from e
